from django.shortcuts import render
from django.http import HttpResponse
from django.template import loader
from django.views.decorators.csrf import csrf_exempt
from docx import Document

from users.functions.functions import handle_uploaded_file
from users.models import User
from users.forms import FirstForm  

def index(request):  
    if request.method == 'POST':  
        student = FirstForm(request.POST, request.FILES)
        name_files = []
        name_files.append(request.FILES)
        full_text = []
        # row_data = []
        if student.is_valid():
            for qty_file in request.FILES.keys():
                if qty_file.startswith('file'): # startswith - true if string value has starts with name file
                    for uploaded_file in request.FILES.getlist(qty_file): # list of files all uploaded
                        document = Document(uploaded_file)
                        for i, table_count in enumerate(document.tables):
                            table = document.tables[i]
                            for x, row in enumerate(table.rows):
                                # for cell in row.cells:
                                #     text = cell.text.strip()
                                text = (cell.text for cell in row.cells)
                                if 'Место проведения испытаний:' in text: # search data by name of first cell
                                    row_data = tuple(text) # кортеж
                                    full_text.append(row_data)
                            text_document = full_text
                    context = {'array_name_files': name_files,
                               'text_document': text_document}
                    return render(request, 'index.html', context)
            # return HttpResponse("File uploaded successfuly")  
    else:  
        student = FirstForm()  
        return render(request,"index.html",{'form':student})


@csrf_exempt
def readDocx(request):
    if request.method == 'POST':
        files = request.POST, request.FILES
        text = request.FILES['file']
        if 'docfile' in request.FILES: 
            extracted_text = [] 
            for file in files: 
                try: 
                    doc = Document(file) 
                    text = ''.join([paragraph.text for paragraph in doc.paragraphs]) 
                    extracted_text.append(text) 
                except Exception as e: 
                    extracted_text.append(f"Ошибка обработки файла: {e}") 
                    context = {'success': 'Files uploaded', 'extracted_text': extracted_text} 
                return render(request, 'formupload.html', context)
        else:
            # document = Document(file)
            # document = Document('C:/Users/User/Desktop/Testing/newProtocols/- JUNFENG T/JUNFENG T80N грузовой автомобиль vin LGFTRA1LXRH113140.docx')
            # full_text = []
            # for par in document.paragraphs:
            #     full_text.append(par.text)
            # text = ' '.join(full_text)
            context = {'error': 'Document not uploaded', 'name_document': files, 
                        'text_document': text}
            return render(request, 'formupload.html', context)
    else:
        context = {'error': 'Document not upload error'}
        return render(request, 'formupload.html', context)
    

@csrf_exempt
def showform(request):
    template = loader.get_template('formupload.html')
    return HttpResponse(template.render())


def users(request):
    allusers = User.objects.all().values()
    template = loader.get_template('all_users.html')
    context = {
        'allusers': allusers,
    }
    return HttpResponse(template.render(context, request))